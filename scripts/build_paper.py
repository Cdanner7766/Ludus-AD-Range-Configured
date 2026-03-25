#!/usr/bin/env python3
"""
Regenerate CIS405_Final_Paper_Danner.docx with revised content.
Run from the repo root: python3 scripts/build_paper.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── constants ─────────────────────────────────────────────────────────────────
FONT_NAME   = "Times New Roman"
FONT_SIZE   = Pt(12)
DBL_SPACE   = Emu(304800)   # 24 pt = double spacing at 12 pt
HALF_INCH   = Emu(457200)   # 0.5-inch first-line / hanging indent

# ── document setup ────────────────────────────────────────────────────────────
doc = Document()

sec = doc.sections[0]
sec.page_width  = Emu(7772400)   # 8.5 in
sec.page_height = Emu(10058400)  # 11 in
for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, attr, Inches(1))

# Remove default paragraph spacing added by some Word styles
style = doc.styles["Normal"]
style.font.name      = FONT_NAME
style.font.size      = FONT_SIZE
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after  = Pt(0)
pf.line_spacing = DBL_SPACE


# ── helpers ───────────────────────────────────────────────────────────────────

def para(text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
         first_line_indent=None, hanging_indent=None):
    """Add a paragraph and return it."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = DBL_SPACE
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if hanging_indent is not None:
        pf.left_indent        = hanging_indent
        pf.first_line_indent  = -hanging_indent
    if text:
        run = p.add_run(text)
        run.bold      = bold
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE
    return p


def heading(text):
    """Centered bold section heading."""
    para(text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


def body(text):
    """Indented body paragraph."""
    para(text, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=HALF_INCH)


def blank():
    """Empty double-spaced line."""
    para()


def ref(text):
    """APA hanging-indent reference entry."""
    para(text, hanging_indent=HALF_INCH)


def add_hyperlink(paragraph, url, display_text):
    """Insert a clickable hyperlink run into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = display_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    # restore font after hyperlink
    run = paragraph.add_run()
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    return hyperlink


# ── PAGE 1: TITLE PAGE ────────────────────────────────────────────────────────

blank()
blank()
blank()
blank()
para("Cloud Deployable Collegiate Cyber Defense Competition Practice Range",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
blank()
para("Charles Danner",        align=WD_ALIGN_PARAGRAPH.CENTER)
para("CIS 405: Cloud Computing", align=WD_ALIGN_PARAGRAPH.CENTER)
para("March 25, 2026",        align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ── BODY ──────────────────────────────────────────────────────────────────────

heading("Cloud Deployable Collegiate Cyber Defense Competition Practice Range")
blank()

# ── 1 ─────────────────────────────────────────────────────────────────────────
heading("Original Project Definition")
body(
    "The project began as an attempt to manually build a Collegiate Cyber "
    "Defense Competition (CCDC) practice environment. The initial plan was to "
    "create and configure virtual machines by hand, run scripts to introduce "
    "security vulnerabilities, and take snapshots so the environment could be "
    "reset between practice sessions. When copies of virtual machines from "
    "previous CCDC competitions were found online, those were tested as a "
    "shortcut (Western Regional CCDC, n.d.). However, none of them came with "
    "documentation, the virtual hardware had to be reconfigured manually for "
    "each machine, and the router's network address translation could not be "
    "made to work correctly. A CCDC environment designer recommended Ludus as "
    "a better approach. From that point the project shifted to building a "
    "fully automated range using Proxmox, Ludus, Ansible, and Claude Code."
)

# ── 2 ─────────────────────────────────────────────────────────────────────────
heading("Changes Made During the Project")
body(
    "Following the Ludus tutorials and writing configuration files and Ansible "
    "roles by hand proved to be extremely slow and had a steep learning curve. "
    "After initial progress stalled, the workflow shifted to using Claude Code "
    "to generate the range configuration and Ansible roles. This change made "
    "it possible to produce a working environment far more quickly. The scope "
    "also grew beyond the original proposal. A custom scoring engine was added "
    "because real CCDC competitions use an automated system that checks whether "
    "services are still running and awards points for uptime (Sshell, 2025). "
    "A referee toggle was added to the scoring dashboard to hide diagnostic "
    "details from the blue team during practice, which mirrors how real "
    "competitions work."
)

# ── 3 ─────────────────────────────────────────────────────────────────────────
heading("Final Scope")
body(
    "By the end of the project the following was completed: Ludus installed "
    "and running on a local Proxmox host; a Windows Active Directory domain "
    "with a domain controller and a Windows 11 workstation; five Linux "
    "service machines running web, database, file sharing, mail, and FTP "
    "services; a Kali Linux attacker machine loaded with a full penetration "
    "testing toolset; a custom Python based scoring engine that checks eleven "
    "services every thirty seconds; a live web dashboard showing service "
    "status and uptime history; and a referee toggle to hide diagnostic "
    "details from blue team players. All nine virtual machines are spread "
    "across two isolated networks and configured entirely through Ansible "
    "roles. Because Ludus supports deployment to any compatible host, this "
    "same environment can be redeployed to Azure, bare metal servers, Google "
    "Cloud Platform, Hyper V, Proxmox, and VMware Fusion without changes to "
    "the configuration."
)

# ── 4 ─────────────────────────────────────────────────────────────────────────
heading("Key Decisions and Why")
body(
    "The most important early decision was to stop building the environment by "
    "hand and adopt Ludus as the base platform. The manual approach required "
    "configuring every virtual machine individually, managing templates and "
    "snapshots without automation, and debugging networking problems that had "
    "no clear solution. Ludus replaced all of that with a single configuration "
    "file and automated provisioning. This is the same infrastructure as code "
    "model that cloud platforms like Azure and Google Cloud use. Describing "
    "the entire environment in a file and letting a tool build it "
    "automatically is what makes the range portable across platforms."
)
body(
    "The decision to use Claude Code rather than write Ansible roles manually "
    "was equally important. Manual role writing stalled progress because of "
    "the large number of tasks needed to configure Active Directory, Linux "
    "services, and network settings correctly. Claude Code generated working "
    "role structure quickly, suggested fixes when tasks failed, and made it "
    "possible to finish all ten roles on schedule. Running a full system "
    "upgrade before installing tools on the Kali Linux machine was a smaller "
    "but necessary decision made after package dependency errors blocked "
    "installation."
)

# ── 5 ─────────────────────────────────────────────────────────────────────────
heading("Tools and Software Used")
body(
    "Proxmox Virtual Environment is the hypervisor that runs all of the "
    "virtual machines in the range. It creates and manages the VMs and "
    "provides the virtual networking that keeps the corporate network and the "
    "attacker network separated. Proxmox is also the platform that Ludus "
    "runs on, making it the foundation for everything else in the project."
)
body(
    "Ludus Cyber Range sits on top of Proxmox and automates the creation of "
    "the entire range from a single configuration file. Ludus creates each "
    "machine from a template, places it on the correct virtual network, and "
    "calls Ansible to finish configuration. Ludus follows the same "
    "infrastructure as code model used by cloud providers and supports "
    "deployment to Azure, bare metal servers, Google Cloud Platform, Hyper V, "
    "Proxmox, and VMware Fusion. This makes the range portable to any "
    "platform where Ludus can run (Ludus, n.d.)."
)
body(
    "Ansible configures each virtual machine after it is created by Ludus. "
    "The ten Ansible roles written for this project cover every machine in "
    "the range, from the domain controller to the FTP server. Each role also "
    "applies intentional security weaknesses so the blue team has realistic "
    "problems to find and fix. Ansible makes the entire configuration "
    "repeatable because it always runs the same steps in the same order from "
    "files stored in the repository."
)
body(
    "Claude Code was used to write and debug the Ansible roles and the Ludus "
    "configuration file. Writing roles for Windows Active Directory and "
    "multiple Linux services involves many tasks, module options, and edge "
    "cases that are difficult to get right manually. Claude Code generated "
    "role structure, suggested fixes when tasks failed, and translated "
    "competition service requirements into working Ansible code. It made "
    "completing all ten roles within the project timeline possible."
)

# ── 6 ─────────────────────────────────────────────────────────────────────────
heading("Diagrams")
body(
    "Figure 1 below shows the full network layout of the range. It includes "
    "the Debian router that connects the WAN, VLAN 10, and VLAN 99, along "
    "with each virtual machine labeled with its IP address, operating system, "
    "purpose, and key ports."
)

diagram_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "network_diagram.png"
)
if os.path.exists(diagram_path):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(0)
    p_img.paragraph_format.space_after  = Pt(0)
    p_img.paragraph_format.line_spacing = DBL_SPACE
    p_img.add_run().add_picture(diagram_path, width=Inches(6.0))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(0)
    p_cap.paragraph_format.space_after  = Pt(0)
    p_cap.paragraph_format.line_spacing = DBL_SPACE
    cap_run = p_cap.add_run("Figure 1. CCDC Practice Range Network Diagram")
    cap_run.italic     = True
    cap_run.font.name  = FONT_NAME
    cap_run.font.size  = FONT_SIZE

# ── 7 ─────────────────────────────────────────────────────────────────────────
heading("Code and Configuration Steps")
body(
    "All configuration files live in the project repository. The main range "
    "configuration file defines every virtual machine, its VLAN, its IP "
    "address, and which Ansible roles to run on it. The ten Ansible roles are "
    "in the roles folder, one subfolder per machine. Each role subfolder "
    "contains a tasks file with the configuration steps, a defaults file with "
    "variable defaults, and a templates folder with config files for services "
    "like Apache, Postfix, and the scoring engine. No configuration is done "
    "by hand; every setting is driven by these files."
)

# ── 8 ─────────────────────────────────────────────────────────────────────────
heading("Installation and Deployment Instructions")

def r(paragraph, text, bold=False):
    run = paragraph.add_run(text)
    run.bold      = bold
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    return run

# Paragraph 1: point to the two resources with embedded hyperlinks
p1 = doc.add_paragraph()
p1.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.LEFT
p1.paragraph_format.space_before      = Pt(0)
p1.paragraph_format.space_after       = Pt(0)
p1.paragraph_format.line_spacing      = DBL_SPACE
p1.paragraph_format.first_line_indent = HALF_INCH
r(p1, "Full deployment instructions for this range are available in ")
add_hyperlink(p1,
              "https://github.com/Cdanner7766/Ludus-AD-Range-Configured/blob/main/SETUP.md",
              "SETUP.md")
r(p1, " in the project repository. Teams that have not yet installed Ludus "
      "should begin with the Ludus quick start guide at ")
add_hyperlink(p1, "https://docs.ludus.cloud/docs/quick-start/",
              "https://docs.ludus.cloud/docs/quick-start/")
r(p1, ".")

# Paragraph 2: the high-level steps as clean prose
body(
    "At a high level, deployment involves five steps. First, install Ludus on "
    "a Proxmox host and build the five required virtual machine templates "
    "(Ludus, n.d.). Second, clone the project repository onto the Ludus host. "
    "Third, register the ten Ansible roles from the repository with Ludus. "
    "Fourth, apply the range configuration file to define the virtual machines "
    "and networks. Fifth, run the range deploy command and allow Ansible to "
    "configure each machine automatically. Once deployment finishes, the "
    "scoring engine starts on its own as a system service and is reachable "
    "on port 8080 from a web browser."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────

blank()
heading("References")
blank()

ref("Ludus. (n.d.). Introduction to Ludus. https://docs.ludus.cloud/docs/intro/")
ref("Ludus. (n.d.). Quick start guide. https://docs.ludus.cloud/docs/quick-start/")
ref("Sshell. (2025). Red teaming at national CCDC 2025. https://www.sshell.co/red-teaming-at-national-ccdc-2025/")
ref("Western Regional CCDC. (n.d.). WRCCDC VM image archive. https://archive.wrccdc.org/images/")
ref("Winterknight. (n.d.). How to win CCDC. https://www.winterknight.net/how-to-win-ccdc/")

# ── save ──────────────────────────────────────────────────────────────────────

out = "CIS405_Final_Paper_Danner.docx"
doc.save(out)
print(f"Saved: {out}")
